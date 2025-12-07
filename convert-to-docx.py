#!/usr/bin/env python3
"""Convert video-scripts-compilation.md to Word document."""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to Word document with proper formatting."""

    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Process content line by line
    lines = content.split('\n')
    in_code_block = False
    code_content = []

    for line in lines:
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block - add accumulated content
                if code_content:
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    for code_line in code_content:
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Skip empty lines (add spacing)
        if not line.strip():
            continue

        # Handle headers
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith('---'):
            # Add a subtle separator
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
            apply_inline_formatting(p)
        elif line.startswith('| '):
            # Table row - simplified handling
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            apply_inline_formatting(p, line)

    # Save document
    doc.save(docx_path)
    print(f"Created: {docx_path}")

def apply_inline_formatting(paragraph, text=None):
    """Apply bold/italic formatting from markdown syntax."""
    if text is None:
        text = paragraph.text
        paragraph.clear()

    # Pattern for **bold** and *italic*
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    md_path = r'C:\Users\james\OneDrive\Documents2\GitHub\business-of-marketing-in-sport\video-scripts-compilation.md'
    docx_path = r'C:\Users\james\OneDrive\Documents2\GitHub\business-of-marketing-in-sport\video-scripts-compilation.docx'
    convert_md_to_docx(md_path, docx_path)
